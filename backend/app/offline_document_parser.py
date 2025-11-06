import os
import re
import hashlib
import json
import mimetypes
from datetime import datetime
from typing import Dict, List, Optional, Set, Any
from collections import defaultdict

class ForensicDocumentExtractor:
    def __init__(self):
        self.mode = 'EXTRACT'  # EXTRACT or ANSWER
        self.extracted_data = None
        self.response_cache = {}

    def extract_document(self, file_path: str) -> Dict[str, Any]:
        """EXTRACT mode: Forensic-level document extraction"""
        self.mode = 'EXTRACT'
        
        if not os.path.exists(file_path):
            return self._create_error_response(file_path, "File not found")
        
        try:
            # Get file metadata
            file_stats = os.stat(file_path)
            filename = os.path.basename(file_path)
            mime_type, _ = mimetypes.guess_type(file_path)
            
            # Extract text content
            full_text = self._extract_text_content(file_path)
            
            # Build forensic extraction JSON
            extraction = {
                'file': {
                    'filename': filename,
                    'mime_type': mime_type or 'application/octet-stream',
                    'size_bytes': file_stats.st_size,
                    'created_at': datetime.fromtimestamp(file_stats.st_ctime).isoformat(),
                    'modified_at': datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
                    'pages': self._estimate_pages(full_text)
                },
                'languages': [{'page': None, 'language': 'en', 'confidence': 0.95}],
                'metadata': self._extract_metadata(file_path, full_text),
                'full_text': full_text,
                'blocks': self._extract_blocks(full_text),
                'tables': self._extract_tables(full_text),
                'forms': self._extract_forms(full_text),
                'images': [],
                'attachments': [],
                'chunks': self._create_chunks(full_text),
                'entities': self._extract_entities(full_text),
                'extraction_remarks': []
            }
            
            self.extracted_data = extraction
            return extraction
            
        except Exception as e:
            return self._create_error_response(file_path, str(e))
    
    def answer_query(self, query: str) -> str:
        """ANSWER mode: Query against extracted data"""
        self.mode = 'ANSWER'
        
        if not self.extracted_data:
            return "ANSWER: No document extracted\nSOURCES: None\nEXCERPTS: None\nCONFIDENCE: Low"
        
        # Search in extracted content
        results = self._search_extracted_data(query)
        
        if not results:
            related = self._find_related_content(query)
            if related:
                return f"ANSWER: Not found in document\nSOURCES: Related content found\nEXCERPTS: {'; '.join(related[:3])}\nCONFIDENCE: Low"
            return "ANSWER: Not found in document\nSOURCES: None\nEXCERPTS: None\nCONFIDENCE: Low"
        
        # Format response
        answer = results['answer']
        sources = '\n'.join([f"• {s}" for s in results['sources']])
        excerpts = '\n'.join([f"• {e}" for e in results['excerpts']])
        confidence = results['confidence']
        
        return f"ANSWER: {answer}\nSOURCES:\n{sources}\nEXCERPTS:\n{excerpts}\nCONFIDENCE: {confidence}"
    
    def _extract_text_content(self, file_path: str) -> str:
        """Extract text from various file formats"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            elif file_ext == '.pdf':
                return self._extract_pdf_text(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_docx_text(file_path)
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
        except Exception:
            return ""
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF"""
        try:
            import fitz
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                return text
            except ImportError:
                return ""
        except Exception:
            return ""

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX"""
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return text
        except ImportError:
            try:
                import docx2txt
                return docx2txt.process(file_path)
            except ImportError:
                return ""
        except Exception:
            return ""
    
    def _create_error_response(self, file_path: str, error: str) -> Dict[str, Any]:
        """Create error response in extraction format"""
        return {
            'file': {'filename': os.path.basename(file_path), 'mime_type': None, 'size_bytes': 0, 'created_at': None, 'modified_at': None, 'pages': None},
            'languages': [],
            'metadata': {'author': None, 'title': None, 'producer': None, 'custom': {}},
            'full_text': '',
            'blocks': [],
            'tables': [],
            'forms': [],
            'images': [],
            'attachments': [],
            'chunks': [],
            'entities': [],
            'extraction_remarks': [{'issue': 'extraction_error', 'page': None, 'note': error}]
        }
    
    def _estimate_pages(self, text: str) -> int:
        """Estimate page count"""
        return max(1, len(text) // 3000) if text else None
    
    def _extract_metadata(self, file_path: str, text: str) -> Dict[str, Any]:
        """Extract document metadata"""
        metadata = {'author': None, 'title': None, 'producer': None, 'custom': {}}
        
        # Extract from text patterns
        author_match = re.search(r'(?:Author|Doctor|Physician)[:\s]+([^\n]+)', text, re.IGNORECASE)
        if author_match:
            metadata['author'] = author_match.group(1).strip()
        
        title_match = re.search(r'(?:Title|Subject)[:\s]+([^\n]+)', text, re.IGNORECASE)
        if title_match:
            metadata['title'] = title_match.group(1).strip()
        
        return metadata
    
    def _extract_blocks(self, text: str) -> List[Dict[str, Any]]:
        """Extract text blocks with provenance"""
        blocks = []
        lines = text.split('\n')
        offset = 0
        
        for i, line in enumerate(lines):
            if line.strip():
                block_type = 'heading' if line.isupper() and len(line) > 5 else 'paragraph'
                blocks.append({
                    'id': f'b{i+1}',
                    'page': None,
                    'block_type': block_type,
                    'text': line.strip(),
                    'bbox': {'x': 0, 'y': i, 'w': len(line), 'h': 1},
                    'start_offset': offset,
                    'end_offset': offset + len(line),
                    'confidence': 0.95
                })
            offset += len(line) + 1
        
        return blocks
    
    def _extract_tables(self, text: str) -> List[Dict[str, Any]]:
        """Extract tables from text"""
        tables = []
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            if '|' in line or '\t' in line:
                # Simple table detection
                cells = re.split(r'[|\t]+', line.strip())
                if len(cells) > 1:
                    tables.append({
                        'table_id': f't{len(tables)+1}',
                        'page': None,
                        'headers': cells if i == 0 else [],
                        'rows': [cells],
                        'csv_snippet': ','.join(cells[:3]),
                        'confidence': 0.8
                    })
        
        return tables
    
    def _extract_forms(self, text: str) -> List[Dict[str, Any]]:
        """Extract form key-value pairs"""
        forms = []
        lines = text.split('\n')
        
        for line in lines:
            if ':' in line:
                parts = line.split(':', 1)
                if len(parts) == 2:
                    key = parts[0].strip()
                    value = parts[1].strip()
                    if key and value:
                        forms.append({
                            'key': key,
                            'value': value,
                            'page': None,
                            'confidence': 0.9
                        })
        
        return forms

    def _create_chunks(self, text: str) -> List[Dict[str, Any]]:
        """Create text chunks"""
        chunks = []
        chunk_size = 2000
        
        for i in range(0, len(text), chunk_size):
            chunk_text = text[i:i+chunk_size]
            chunks.append({
                'chunk_id': f'c{len(chunks)+1}',
                'start_offset': i,
                'end_offset': min(i+chunk_size, len(text)),
                'text': chunk_text,
                'page_range': [None, None]
            })
        
        return chunks
    
    def _extract_entities(self, text: str) -> List[Dict[str, Any]]:
        """Extract named entities"""
        entities = []
        
        # Simple pattern-based entity extraction
        person_pattern = r'Dr\. [A-Z][a-z]+ [A-Z][a-z]+'
        for match in re.finditer(person_pattern, text):
            entities.append({
                'text': match.group(),
                'type': 'PERSON',
                'page': None,
                'confidence': 0.9
            })
        
        return entities
    
    def _search_extracted_data(self, query: str) -> Dict[str, Any]:
        """Search within extracted data"""
        query_lower = query.lower()
        results = {'answer': '', 'sources': [], 'excerpts': [], 'confidence': 'Low'}
        
        # Search in full text
        if query_lower in self.extracted_data['full_text'].lower():
            # Find matching blocks
            for block in self.extracted_data['blocks']:
                if query_lower in block['text'].lower():
                    results['sources'].append(f"block {block['id']}")
                    results['excerpts'].append(block['text'][:100])
            
            # Search forms
            for form in self.extracted_data['forms']:
                if query_lower in form['key'].lower() or query_lower in form['value'].lower():
                    results['sources'].append(f"form {form['key']}")
                    results['excerpts'].append(f"{form['key']}: {form['value']}")
            
            if results['sources']:
                results['answer'] = results['excerpts'][0] if results['excerpts'] else 'Found in document'
                results['confidence'] = 'High' if len(results['sources']) > 1 else 'Medium'
        
        return results if results['sources'] else None
    
    def _find_related_content(self, query: str) -> List[str]:
        """Find related content for failed searches"""
        query_words = set(query.lower().split())
        related = []
        
        for block in self.extracted_data['blocks']:
            block_words = set(block['text'].lower().split())
            if query_words.intersection(block_words):
                related.append(block['text'][:100])
        
        return related
    
    def extract_document_from_text(self, text: str):
        """Extract from text content directly for testing"""
        self.extracted_data = {
            'file': {'filename': 'text_input', 'mime_type': 'text/plain', 'size_bytes': len(text), 'created_at': None, 'modified_at': None, 'pages': None},
            'languages': [{'page': None, 'language': 'en', 'confidence': 0.95}],
            'metadata': {'author': None, 'title': None, 'producer': None, 'custom': {}},
            'full_text': text,
            'blocks': self._extract_blocks(text),
            'tables': self._extract_tables(text),
            'forms': self._extract_forms(text),
            'images': [],
            'attachments': [],
            'chunks': self._create_chunks(text),
            'entities': self._extract_entities(text),
            'extraction_remarks': []
        }

# Backward compatibility alias
OfflineDocumentParser = ForensicDocumentExtractor

class OfflineDocumentParser(ForensicDocumentExtractor):
    """Legacy class for backward compatibility"""
    def parse_document(self, text: str) -> Dict[str, Dict[str, str]]:
        """Legacy method that returns old format"""
        # Extract using forensic method
        if text:
            self.extract_document_from_text(text)
        
        # Convert to legacy format
        if hasattr(self, 'extracted_data') and self.extracted_data:
            return self._convert_forensic_to_legacy_format(self.extracted_data)
        return {}
    
    def extract_text(self, file_path: str) -> str:
        """Legacy text extraction method"""
        return self._extract_text_content(file_path)
    
    def extract_document_from_text(self, text: str):
        """Extract from text content directly"""
        # Create minimal extraction data
        self.extracted_data = {
            'file': {'filename': 'text_input', 'mime_type': 'text/plain', 'size_bytes': len(text), 'created_at': None, 'modified_at': None, 'pages': None},
            'languages': [{'page': None, 'language': 'en', 'confidence': 0.95}],
            'metadata': {'author': None, 'title': None, 'producer': None, 'custom': {}},
            'full_text': text,
            'blocks': self._extract_blocks(text),
            'tables': self._extract_tables(text),
            'forms': self._extract_forms(text),
            'images': [],
            'attachments': [],
            'chunks': self._create_chunks(text),
            'entities': self._extract_entities(text),
            'extraction_remarks': []
        }
    
    def _convert_forensic_to_legacy_format(self, forensic_data: Dict[str, Any]) -> Dict[str, Dict[str, str]]:
        """Convert forensic format to legacy structured format"""
        legacy_format = {}
        
        # Patient information
        patient_info = {}
        for form in forensic_data.get('forms', []):
            key = form.get('key', '').lower()
            value = form.get('value', '')
            if 'name' in key:
                patient_info['patient_name'] = value
            elif 'age' in key:
                patient_info['age'] = value
            elif 'gender' in key:
                patient_info['gender'] = value
        
        if patient_info:
            legacy_format['patient_information'] = patient_info
        
        # Extract other sections from text patterns
        text = forensic_data.get('full_text', '')
        legacy_format.update(self._extract_legacy_sections(text))
        
        return legacy_format
    
    def _extract_legacy_sections(self, text: str) -> Dict[str, Dict[str, str]]:
        """Extract sections in legacy format"""
        sections = {}
        
        # Diagnosis
        diag_match = re.search(r'DIAGNOSIS[:\s]*(.*?)(?=CLINICAL|INVESTIGATIONS|TREATMENT|$)', text, re.IGNORECASE | re.DOTALL)
        if diag_match:
            sections['diagnosis'] = {'diagnosis': diag_match.group(1).strip()}
        
        # Medications
        med_match = re.search(r'DISCHARGE MEDICATIONS[:\s]*(.*?)(?=ADVICE|FOLLOW|$)', text, re.IGNORECASE | re.DOTALL)
        if med_match:
            sections['discharge_medications'] = {'medications': med_match.group(1).strip()}
        
        # Clinical summary
        summary_match = re.search(r'CLINICAL SUMMARY[:\s]*(.*?)(?=INVESTIGATIONS|TREATMENT|$)', text, re.IGNORECASE | re.DOTALL)
        if summary_match:
            sections['clinical_summary'] = {'summary': summary_match.group(1).strip()}
        
        return sections

    def _extract_all_patient_info(self, text: str) -> Dict[str, str]:
        """Extract ALL patient information"""
        info = {}
        lines = text.split('\n')
        
        # Extract all patient fields
        for line in lines:
            line = line.strip()
            if ':' in line:
                key, value = line.split(':', 1)
                key = key.strip().lower()
                value = value.strip()
                
                if 'patient name' in key or key == 'name':
                    info['patient_name'] = value
                elif 'age' in key:
                    age_match = re.search(r'(\d+)', value)
                    if age_match:
                        info['age'] = age_match.group(1)
                elif 'gender' in key or 'sex' in key:
                    info['gender'] = value
                elif 'mrn' in key or 'medical record' in key or 'hospital number' in key:
                    info['hospital_number'] = value
                elif 'date of birth' in key or 'dob' in key:
                    info['date_of_birth'] = value
                elif 'admission' in key:
                    info['admission_date'] = value
                elif 'discharge' in key and 'date' in key:
                    info['discharge_date'] = value
                elif 'physician' in key or 'doctor' in key:
                    info['attending_physician'] = value
                elif 'unit' in key or 'room' in key:
                    info['unit_room'] = value
        
        # Handle "Age / Gender: 63 / Female" format
        for line in lines:
            if 'Age / Gender:' in line:
                parts = line.split(':', 1)[1].strip().split('/')
                if len(parts) >= 2:
                    info['age'] = parts[0].strip()
                    info['gender'] = parts[1].strip()
        
        return info

    def _extract_all_diagnosis(self, text: str) -> Dict[str, str]:
        """Extract ONLY diagnosis list"""
        diagnosis = {}
        lines = text.split('\n')
        
        # Find diagnosis section
        for i, line in enumerate(lines):
            if any(keyword in line.upper() for keyword in ['PRIMARY DIAGNOSIS', 'DIAGNOSIS', 'DIAGNOSES']):
                # Extract only numbered diagnosis items
                diagnosis_lines = []
                j = i + 1
                while j < len(lines):
                    next_line = lines[j].strip()
                    if not next_line:
                        j += 1
                        continue
                    # Stop at next major section or clinical summary
                    if (self._is_major_section(next_line) or 
                        'CLINICAL SUMMARY' in next_line.upper()):
                        break
                    # Only include numbered diagnosis lines
                    if (next_line.startswith(('1.', '2.', '3.', '4.', '5.')) or
                        (len(diagnosis_lines) == 0 and any(word in next_line.lower() for word in ['diabetes', 'hypertension', 'mellitus']))):
                        diagnosis_lines.append(next_line)
                    j += 1
                
                if diagnosis_lines:
                    diagnosis['diagnosis'] = '\n'.join(diagnosis_lines)
                    break
        
        return diagnosis

    def _extract_all_clinical(self, text: str) -> Dict[str, str]:
        """Extract ALL clinical information including summary"""
        clinical = {}
        
        # Check if entire text is a summary (no section headers)
        if not any(header in text.upper() for header in ['DIAGNOSIS', 'MEDICATIONS', 'TREATMENT', 'INVESTIGATIONS']):
            # Treat entire text as summary
            clinical['summary'] = text.strip()
            return clinical
        
        lines = text.split('\n')
        
        # Find clinical sections
        clinical_keywords = ['HOSPITAL COURSE', 'CLINICAL SUMMARY', 'CLINICAL PRESENTATION', 'HISTORY', 'SUMMARY']
        
        for i, line in enumerate(lines):
            if any(keyword in line.upper() for keyword in clinical_keywords):
                clinical_lines = []
                j = i + 1
                while j < len(lines):
                    next_line = lines[j].strip()
                    if not next_line:
                        j += 1
                        continue
                    if self._is_major_section(next_line):
                        break
                    clinical_lines.append(next_line)
                    j += 1
                
                if clinical_lines:
                    clinical['summary'] = '\n'.join(clinical_lines)
                    break
        
        # Extract vital signs from anywhere in text
        vital_patterns = {
            'blood_pressure': r'(?:Blood pressure|BP)\s*[:\-]?\s*(\d+/\d+)',
            'heart_rate': r'(?:Heart rate|HR|Pulse)\s*[:\-]?\s*(\d+)',
            'temperature': r'(?:Temperature|Temp)\s*[:\-]?\s*(\d+\.?\d*)',
            'respiratory_rate': r'(?:Respiratory rate|RR)\s*[:\-]?\s*(\d+)',
            'oxygen_saturation': r'(?:SpO2|O2)\s*[:\-]?\s*(\d+)%?'
        }
        
        for vital, pattern in vital_patterns.items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                clinical[vital] = match.group(1)
        
        return clinical

    def _extract_all_investigations(self, text: str) -> Dict[str, str]:
        """Extract ONLY investigation results"""
        investigations = {}
        lines = text.split('\n')
        
        # Find investigations section
        for i, line in enumerate(lines):
            if 'INVESTIGATIONS' in line.upper():
                inv_lines = []
                j = i + 1
                while j < len(lines):
                    next_line = lines[j].strip()
                    if not next_line:
                        j += 1
                        continue
                    if (self._is_major_section(next_line) or 
                        'VITAL SIGNS' in next_line.upper()):
                        break
                    # Only include test result lines with values
                    if ((':' in next_line and any(word in next_line.lower() for word in ['blood', 'glucose', 'hemoglobin', 'creatinine', 'sodium', 'wbc', 'hba1c'])) or
                        next_line.startswith('-')):
                        inv_lines.append(next_line)
                    j += 1
                
                if inv_lines:
                    investigations['investigations_details'] = '\n'.join(inv_lines)
                    break
        
        return investigations

    def _extract_all_treatment(self, text: str) -> Dict[str, str]:
        """Extract ONLY treatment list"""
        treatment = {}
        lines = text.split('\n')
        
        # Find treatment sections
        treatment_keywords = ['TREATMENT PROVIDED', 'TREATMENT GIVEN', 'TREATMENT']
        
        for i, line in enumerate(lines):
            if any(keyword in line.upper() for keyword in treatment_keywords):
                treatment_lines = []
                j = i + 1
                while j < len(lines):
                    next_line = lines[j].strip()
                    if not next_line:
                        j += 1
                        continue
                    if (self._is_major_section(next_line) or 
                        'DISCHARGE MEDICATIONS' in next_line.upper()):
                        break
                    # Only include numbered treatment lines or relevant treatment
                    if (next_line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.')) or
                        any(word in next_line.lower() for word in ['iv', 'insulin', 'therapy', 'medication', 'fluid', 'treatment'])):
                        treatment_lines.append(next_line)
                    j += 1
                
                if treatment_lines:
                    treatment['treatment'] = '\n'.join(treatment_lines)
                    break
        
        return treatment

    def _extract_all_medications(self, text: str) -> Dict[str, str]:
        """Extract ONLY medication list"""
        medications = {}
        lines = text.split('\n')
        
        # Find medication sections
        for i, line in enumerate(lines):
            if any(keyword in line.upper() for keyword in ['DISCHARGE MEDICATIONS', 'MEDICATIONS', 'PRESCRIPTIONS']):
                med_lines = []
                j = i + 1
                while j < len(lines):
                    next_line = lines[j].strip()
                    if not next_line:
                        j += 1
                        continue
                    if (self._is_major_section(next_line) or 
                        'DISCHARGE ADVICE' in next_line.upper()):
                        break
                    # Only include numbered medication lines or lines with dosage
                    if (next_line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.')) or
                        any(word in next_line.lower() for word in ['mg', 'daily', 'twice', 'tablet', 'capsule'])):
                        med_lines.append(next_line)
                    j += 1
                
                if med_lines:
                    medications['medications'] = '\n'.join(med_lines)
                    break
        
        return medications

    def _extract_all_advice(self, text: str) -> Dict[str, str]:
        """Extract ONLY discharge advice list"""
        advice = {}
        lines = text.split('\n')
        
        # Find advice sections
        advice_keywords = ['DISCHARGE ADVICE', 'INSTRUCTIONS', 'ADVICE']
        
        for i, line in enumerate(lines):
            if any(keyword in line.upper() for keyword in advice_keywords):
                advice_lines = []
                j = i + 1
                while j < len(lines):
                    next_line = lines[j].strip()
                    if not next_line:
                        j += 1
                        continue
                    if (self._is_major_section(next_line) or 
                        'FOLLOW-UP' in next_line.upper()):
                        break
                    # Only include numbered advice lines or relevant advice
                    if (next_line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')) or
                        any(word in next_line.lower() for word in ['follow', 'diet', 'medication', 'monitor', 'avoid', 'take'])):
                        advice_lines.append(next_line)
                    j += 1
                
                if advice_lines:
                    advice['instructions'] = '\n'.join(advice_lines)
                    break
        
        return advice

    def _extract_all_followup(self, text: str) -> Dict[str, str]:
        """Extract ALL follow-up information"""
        followup = {}
        lines = text.split('\n')
        
        # Find follow-up sections
        for i, line in enumerate(lines):
            if any(keyword in line.upper() for keyword in ['FOLLOW-UP', 'FOLLOW UP', 'FOLLOWUP']):
                followup_lines = []
                j = i + 1
                while j < len(lines):
                    next_line = lines[j].strip()
                    if not next_line:
                        j += 1
                        continue
                    if self._is_major_section(next_line):
                        break
                    followup_lines.append(next_line)
                    j += 1
                
                if followup_lines:
                    followup['followup_timing'] = '\n'.join(followup_lines)
                    break
        
        return followup

    def _is_major_section(self, line: str) -> bool:
        """Check if line is a major section header"""
        if not line.isupper() or len(line) < 5:
            return False
        
        major_sections = [
            'PRIMARY DIAGNOSIS', 'DIAGNOSIS', 'HOSPITAL COURSE', 'CLINICAL SUMMARY',
            'INVESTIGATIONS', 'TREATMENT PROVIDED', 'TREATMENT GIVEN', 'DISCHARGE MEDICATIONS',
            'FOLLOW-UP INSTRUCTIONS', 'DISCHARGE ADVICE', 'DISCHARGE SUMMARY', 'PHYSICIAN'
        ]
        
        return any(section in line for section in major_sections)